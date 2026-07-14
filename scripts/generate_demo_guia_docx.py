#!/usr/bin/env python3
"""Genera DEMO_GUIA.docx para examples/demo-repo."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

PROJECT_ROOT = Path(__file__).resolve().parents[1]
OUTPUT = PROJECT_ROOT / "examples" / "demo-repo" / "DEMO_GUIA.docx"


def add_code_block(doc: Document, text: str) -> None:
    for line in text.strip().split("\n"):
        para = doc.add_paragraph(line)
        para.paragraph_format.left_indent = Inches(0.3)
        for run in para.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9)


def main() -> None:
    doc = Document()

    title = doc.add_heading("Guía de demostración — demo-repo", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "TFM: Detección de secretos en código y pipelines\n"
        "Ángela Ramírez Barajas — UCAM, 2026"
    )
    run.italic = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        "Repositorio: examples/demo-repo/ dentro de secret-scan-tfm. "
        "Objetivo: demostrar pre-commit + Gitleaks (bloqueo local) y allowlist de fixtures MOCK_."
    )

    doc.add_heading("Qué contiene el demo", level=1)
    for path, desc in [
        ("app/", "Código de aplicación; credenciales solo vía variables de entorno"),
        ("tests/fixtures/", "Tokens MOCK_* permitidos por .gitleaks.toml"),
        ("leaks/", "Plantilla para provocar un commit bloqueado"),
        (".gitleaks.toml", "Reglas Gitleaks + allowlist documentada"),
        (".pre-commit-config.yaml", "Hook Gitleaks antes de cada commit"),
    ]:
        doc.add_paragraph(f"{path} — {desc}", style="List Bullet")

    doc.add_heading("Requisitos previos", level=1)
    for req in [
        "Git instalado",
        "Python 3.10+",
        "pip install pre-commit",
        "Gitleaks en PATH (https://github.com/gitleaks/gitleaks)",
    ]:
        doc.add_paragraph(req, style="List Bullet")

    doc.add_heading("Paso 1 — Configuración inicial (una vez)", level=1)
    doc.add_paragraph("Abrir terminal (Git Bash o PowerShell) y ejecutar:")
    add_code_block(
        doc,
        r"""cd C:\IA\secret-scan-tfm\examples\demo-repo
git init
pip install pre-commit
pre-commit install""",
    )
    doc.add_paragraph("Salida esperada: pre-commit instalado en .git/hooks/")

    doc.add_heading("Paso 2 — Escaneo correcto (Passed)", level=1)
    doc.add_paragraph("Comprobar que el código versionado no dispara alertas:")
    add_code_block(doc, "pre-commit run gitleaks --all-files")
    doc.add_paragraph(
        "Resultado esperado: Passed. Los MOCK en tests/fixtures/ están en allowlist."
    )

    doc.add_heading("Paso 3 — Demostrar bloqueo de commit (Failed)", level=1)
    doc.add_paragraph("Copiar plantilla con token ficticio e intentar commitear:")
    add_code_block(
        doc,
        """cp leaks/intentional-leak.env.template leaks/demo-block-me.env
git add -f leaks/demo-block-me.env
git commit -m "test: intento de filtrar secreto" """,
    )
    doc.add_paragraph(
        "Resultado esperado: Gitleaks Failed — commit abortado. Muestra archivo y línea."
    )
    doc.add_paragraph("Limpieza: eliminar leaks/demo-block-me.env antes de seguir.")

    doc.add_heading("Paso 4 — Opcional: triaje LLM (secret-triage + Ollama)", level=1)
    doc.add_paragraph(
        "Requiere Ollama con llama3.1:8b en ejecución. No forma parte del pre-commit."
    )
    add_code_block(
        doc,
        """gitleaks detect --source . --report-format json --report-path gitleaks.json --no-git
secret-triage filter --report gitleaks.json --repo-root . --limit 3 -o triaged.json
secret-triage report triaged.json""",
    )

    doc.add_heading("Para la defensa oral", level=1)
    for tip in [
        "Mostrar capturas C04 (bloqueo) y C05 (GitHub Actions) si no hay tiempo para demo en vivo",
        "Mensaje clave: Gitleaks en pre-commit (< 1 s); LLM solo en batch (latencia Ollama)",
        "Duración demo en vivo: 2–3 minutos (Pasos 2 y 3)",
    ]:
        doc.add_paragraph(tip, style="List Bullet")

    doc.add_paragraph()
    note = doc.add_paragraph(
        "Más detalle: README.md y leaks/README.md en esta carpeta. Memoria: capítulo 7.2."
    )
    note.runs[0].italic = True

    doc.save(OUTPUT)
    print(f"Guardado: {OUTPUT}")


if __name__ == "__main__":
    main()
